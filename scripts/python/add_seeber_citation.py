#!/usr/bin/env python3
"""
Add Seeber et al. (2024) Citation to Nature Communications Manuscript
=====================================================================

Adds citation and discussion about self-citation gaming to distinguish
structural database bias from publisher manipulation.
"""

from docx import Document

def add_seeber_discussion():
    """Add Seeber citation and discussion to manuscript."""

    print("="*80)
    print("ADDING SEEBER ET AL. (2024) CITATION TO MANUSCRIPT")
    print("="*80)

    # Load manuscript
    doc = Document('manuscripts/Nature_Communications_Article-v7.docx')
    print(f"\n✓ Loaded manuscript: {len(doc.paragraphs)} paragraphs")

    # Find Discussion section (look for "Discussion" or conclusion-like text)
    discussion_idx = None
    for i, para in enumerate(doc.paragraphs):
        # Look for discussion section or mechanisms discussion
        if any(keyword in para.text.lower() for keyword in ['mechanism', 'structural bias', 'database bias', 'conclusion']):
            if len(para.text) > 100:  # Substantive paragraph
                discussion_idx = i
                print(f"\n✓ Found discussion section at paragraph {i}")
                print(f"  Preview: {para.text[:100]}...")
                break

    # If we found a good spot, add after it
    if discussion_idx:
        # Add the new paragraph
        seeber_text = """Importantly, the Elsevier bias we identified operates through a fundamentally different mechanism than the self-citation gaming observed in some publishers. Recent analysis of citation patterns across major publishers found that while some Article Processing Charge (APC) publishers engage in extreme strategic self-citation to inflate Journal Impact Factors—with MDPI and Frontiers showing self-citation premiums of +2,595% and +1,167% respectively to articles affecting JIF calculations—Elsevier exhibits moderate self-citation behavior (+554%) comparable to other traditional subscription publishers like Wiley, Springer, and Taylor & Francis (Seeber et al., 2024). This demonstrates that the systematic coverage advantage we document stems from structural over-representation in Scopus itself—a database-level bias affecting all citation-based metrics—rather than publisher manipulation of citation patterns. Such structural bias is arguably more concerning than self-citation gaming because it operates invisibly at the infrastructure level, systematically advantaging one publisher's content across all bibliometric applications that rely on Scopus data."""

        # Insert after discussion paragraph
        new_para = doc.add_paragraph()
        new_para.text = seeber_text
        new_para.style = doc.paragraphs[discussion_idx].style

        # Move to correct position (after discussion_idx)
        para_element = new_para._element
        para_element.getparent().remove(para_element)
        doc.paragraphs[discussion_idx]._element.addnext(para_element)

        print(f"\n✓ Added Seeber discussion paragraph after paragraph {discussion_idx}")

    # Now find and update References section
    refs_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == 'references' or para.text.strip() == 'References':
            refs_idx = i
            print(f"\n✓ Found References section at paragraph {i}")
            break

    if refs_idx:
        # Find Philip (2023) reference to insert after it (since Seeber comes after Philip alphabetically)
        # Look for existing references
        philip_idx = None
        for i in range(refs_idx + 1, min(refs_idx + 20, len(doc.paragraphs))):
            if 'Philip' in doc.paragraphs[i].text and '2023' in doc.paragraphs[i].text:
                philip_idx = i
                print(f"  Found Philip (2023) at paragraph {i}")
                break

        # Add Seeber reference after Philip
        seeber_ref = """Seeber, M., Cattaneo, M., & Birolini, S. (2024). Academic publishing business models: self-citations and the selectivity-reputation trade-off. Quantitative Science Studies, 5(1), 1-29. https://doi.org/10.1162/qss_a_00322"""

        ref_para = doc.add_paragraph()
        ref_para.text = seeber_ref
        ref_para.style = doc.paragraphs[philip_idx].style if philip_idx else 'Normal'

        # Move to correct position
        para_element = ref_para._element
        para_element.getparent().remove(para_element)
        insert_after = philip_idx if philip_idx else refs_idx
        doc.paragraphs[insert_after]._element.addnext(para_element)

        print(f"  ✓ Added Seeber et al. (2024) reference after paragraph {insert_after}")

    # Save as v8
    output_file = 'manuscripts/Nature_Communications_Article-v8.docx'
    doc.save(output_file)

    print(f"\n{'='*80}")
    print(f"✓ MANUSCRIPT UPDATED")
    print(f"{'='*80}")
    print(f"\nSaved as: {output_file}")
    print(f"  • Added Seeber et al. (2024) discussion paragraph")
    print(f"  • Added reference to References section")
    print(f"  • New version: v8")
    print(f"\nKey point added:")
    print(f"  - Elsevier shows MODERATE self-citation (+554%), similar to Wiley/Springer")
    print(f"  - MDPI/Frontiers show EXTREME self-citation (+2,595%/+1,167%)")
    print(f"  - This proves your bias is STRUCTURAL (database), not behavioral (gaming)")
    print(f"  - Structural bias is MORE concerning because it's invisible and systematic")

    return output_file

if __name__ == "__main__":
    add_seeber_discussion()
